import os
import re
import pdfplumber
import docx
import pandas as pd
from rich.console import Console
from rich.table import Table
from rich.theme import Theme
from rich import box
import sys
import time
import shutil
import pyfiglet
import requests
import csv
import json
import socket
import base64
from ipaddress import ip_network, ip_address, collapse_addresses
from ipwhois import IPWhois
from concurrent.futures import ThreadPoolExecutor, as_completed

# ---------------------------------------------------------------------
# Initialize rich console with a custom theme
# ---------------------------------------------------------------------
custom_theme = Theme({
    "header": "bold magenta",
    "info": "dim cyan",
    "warning": "bold yellow",
    "error": "bold red",
    "success": "bold green",
})
console = Console(theme=custom_theme)

# --- Helper Functions for Displaying ASCII Art (Optional) ---
def get_terminal_size():
    """Return the size of the terminal window."""
    columns, lines = shutil.get_terminal_size((80, 20))
    return columns, lines

def read_ascii_art(file_path):
    """Read ASCII art from a file."""
    with open(file_path, 'r') as file:
        return file.readlines()

def display_ascii_art(art, columns):
    """Display ASCII art resized to fit the terminal width."""
    for line in art:
        print((line.strip('\n') + ' ' * columns)[:columns])

def banner():
    """Display banner ASCII art."""
    file_path = 'ascii_art.txt'  # Assuming you have an ASCII art file
    if os.path.exists(file_path):
        ascii_art = read_ascii_art(file_path)
        columns, _ = get_terminal_size()
        os.system('cls' if os.name == 'nt' else 'clear')
        display_ascii_art(ascii_art, columns)
        time.sleep(5)  # Reduced wait time
        print("\033c", end="")  # Clear the screen

def title():
    wordart = pyfiglet.figlet_format("ATT", font="dos_rebel")
    print(wordart)

# ---------------------------------------------------------------------
# List of keywords to search for in tables
# ---------------------------------------------------------------------
keywords = [
    'TADIG Code', 'Network Name', 'Network Type', 'Technology', '2G Frequencies', '3G Frequencies',
    '4G Frequencies', '5G Frequencies', 'MSISDN Number Ranges', 'Country Code', 'National Destination Code',
    'SN Range Start', 'SN Range Stop', 'Mobile Country Code', 'Mobile Network Code', 'IMSI to MGT Translation',
    'Does Number Portability Apply?', 'SCCP Gateway Information', 'SCCP Carrier Name', 'DPC Information',
    'XUDT/XUDTS Segmentation Capabilities', 'ANSI Networks SCCP Gateway', 'Subscriber Identity Authentication',
    'Authentication Performed for GSM Subscribers', 'Authentication Performed for GPRS Subscribers', 'Cipher Algorithm',
    'Automatic Roaming Testing', 'AAC (Automatic Answering Circuit) Information',
    'DAAC (Data Automatic Answering Circuit) Information', 'Mobile Application Part (MAP) Information',
    'Application Context Names', 'Inbound Roaming', 'Outbound Roaming', 'MSC/VLR, SGSN versions',
    'MAP Versions', 'MAP Functionality', 'roamingNumberEnquiry', 'subscribeDataMngt', 'CAMEL Information',
    'CAMEL Phases and Versions', 'O-CSI', 'D-CSI', 'MT-SMS-CSI', 'CAP Version', 'CAPv2', 'CAPv3', 'CAPv4',
    'Partial Implementations', 'CAMEL Phase 4 CSIs', 'MG-CSI', 'Network Elements Information', 'Node Types',
    'HLR', 'MSC/VLR', 'SCP', 'SMSC', 'Node IDs and Addresses', 'E.164 Numbers', 'IP Addresses',
    'Global Title Addresses', 'Roamware Dummy Global Titles', 'VIP Information', 'SGSN', 'GGSN VIP IP addresses',
    'Packet Data Services Information', 'APN Identifiers', 'APN Operator Identifier', 'APN Credentials',
    'APN DNS IP Addresses', 'GTP Versions', 'GTPv1', 'SGSN and GGSN Versions', 'Multiple PDP Context Support',
    '2G/3G Data Service Profiles', 'GPRS Information', 'APN List for Testing and Troubleshooting',
    'WAP APN', 'MMS APN Information', 'Pingable and Trace-routable IP Addresses', 'Autonomous System Numbers',
    'ASN', 'GRX Providers', 'Inter-PLMN GSN Backbone IP Address Ranges', 'LTE Roaming Information',
    'IPX Provider Names', 'Primary IP Addresses', 'Secondary IP Addresses', 'Diameter Architecture',
    'EPC Realms for Roaming', 'S6a', 'S6d', 'SMS over NAS Support', 'LTE QoS Profiles', 'Miscellaneous Information',
    'MSRN Structure', 'IMSI Structure', 'SMSC GT Addresses', 'Additional Global Title Addresses', 'Roamware Dummy'
]

# ---------------------------------------------------------------------
# Table Extraction
# ---------------------------------------------------------------------
def table_contains_keywords(table, keywords):
    for row in table:
        for cell in row:
            if cell:
                for keyword in keywords:
                    if keyword.lower() in cell.lower():
                        return True
    return False

def get_table_keyword(table, keywords):
    """
    Return the most relevant keyword found in the table.
    Used to label the CSV with something meaningful.
    """
    keyword_counts = {}
    for row in table:
        for cell in row:
            if cell:
                for keyword in keywords:
                    if keyword.lower() in cell.lower():
                        keyword_counts[keyword] = keyword_counts.get(keyword, 0) + 1
    if keyword_counts:
        sorted_keywords = sorted(keyword_counts.items(), key=lambda x: x[1], reverse=True)
        return sorted_keywords[0][0].replace(' ', '_')  # e.g., "TADIG_Code"
    return 'table'

def extract_tables_from_pdf(file_path, keywords):
    tables = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_tables = page.extract_tables()
            if page_tables:
                for table in page_tables:
                    # Some tables may contain None cells
                    sanitized_table = []
                    for row in table:
                        sanitized_row = [str(cell) if cell else '' for cell in row]
                        sanitized_table.append(sanitized_row)
                    # Now check keywords
                    if table_contains_keywords(sanitized_table, keywords):
                        console.print(
                            f"Found matching table on page {page_num} in [info]{os.path.basename(file_path)}[/info]",
                            style="success")
                        tables.append(sanitized_table)
    return tables

def extract_tables_from_docx(file_path, keywords):
    doc = docx.Document(file_path)
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        if data and table_contains_keywords(data, keywords):
            console.print(
                f"Found matching table in [info]{os.path.basename(file_path)}[/info]",
                style="success")
            tables.append(data)
    return tables

def save_tables(tables, output_directory, base_filename):
    for i, table in enumerate(tables):
        max_columns = max(len(row) for row in table)
        standardized_table = []
        for row in table:
            # Pad or truncate each row to match max_columns
            if len(row) < max_columns:
                row.extend([''] * (max_columns - len(row)))  # <-- Fixed parenthesis here
            elif len(row) > max_columns:
                row = row[:max_columns]
            standardized_table.append(row)

        # Use the first row as the header
        header = standardized_table[0]
        data_rows = standardized_table[1:]
        # Replace empty headers with default names
        header = [f"Column_{j+1}" if not col else col for j, col in enumerate(header)]

        df = pd.DataFrame(data_rows, columns=header)

        # Get a keyword for the table name
        table_keyword = get_table_keyword(standardized_table, keywords)
        table_keyword = re.sub(r'[^\w\-_. ]', '_', table_keyword)

        output_file = os.path.join(output_directory, f"{base_filename}_{table_keyword}_{i+1}.csv")
        df.to_csv(output_file, index=False)
        console.print(f"Saved table to [success]{output_file}[/success]")

def display_tables_from_folder(output_directory, selected_indices=None):
    csv_files = [f for f in os.listdir(output_directory) if f.endswith('.csv')]
    if not csv_files:
        console.print("No tables found in this folder.", style="warning")
        return
    if selected_indices is None:
        selected_indices = range(1, len(csv_files) + 1)
    for idx in selected_indices:
        if 1 <= idx <= len(csv_files):
            csv_file = csv_files[idx - 1]
            df = pd.read_csv(os.path.join(output_directory, csv_file))
            console.print(f"Displaying [info]{csv_file}[/info]:", style="header")
            display_table_with_rich(df)
        else:
            console.print(f"Table number {idx} is out of range.", style="error")

def display_table_with_rich(df):
    table = Table(show_header=True, header_style="bold magenta", box=box.SQUARE)
    for column in df.columns:
        table.add_column(str(column), overflow="fold")
    for _, row in df.iterrows():
        table.add_row(*[str(item) for item in row])
    console.print(table)

def process_file(file_path, directory):
    filename = os.path.basename(file_path)
    console.print(f"Processing file: [info]{filename}[/info]", style="success")
    if filename.lower().endswith(".pdf"):
        tables = extract_tables_from_pdf(file_path, keywords)
    elif filename.lower().endswith(".docx"):
        tables = extract_tables_from_docx(file_path, keywords)
    else:
        console.print(f"Skipping file: [warning]{filename}[/warning] (unsupported format)")
        return

    if tables:
        base_filename = os.path.splitext(filename)[0]
        output_directory = os.path.join(directory, base_filename)
        os.makedirs(output_directory, exist_ok=True)
        save_tables(tables, output_directory, base_filename)
        display_tables_from_folder(output_directory)
    else:
        console.print(f"No matching tables found in [warning]{filename}[/warning]", style="warning")

def list_processed_folders(directory):
    folders = [f for f in os.listdir(directory) if os.path.isdir(os.path.join(directory, f))]
    folders = [f for f in folders if not f.startswith('.')]
    if not folders:
        console.print("No previously ingested folders found.", style="warning")
        return []
    console.print("Previously ingested folders:", style="header")
    for idx, folder in enumerate(folders, 1):
        console.print(f"{idx}. {folder}")
    return folders

def parse_user_selection(user_input, max_value):
    selections = set()
    parts = user_input.split(',')
    for part in parts:
        part = part.strip()
        if '-' in part:
            try:
                start_str, end_str = part.split('-')
                start = int(start_str)
                end = int(end_str)
                if start <= end:
                    for i in range(start, end + 1):
                        if 1 <= i <= max_value:
                            selections.add(i)
                else:
                    console.print(f"Invalid range '{part}'.", style="error")
            except ValueError:
                console.print(f"Invalid range '{part}'.", style="error")
        else:
            try:
                i = int(part)
                if 1 <= i <= max_value:
                    selections.add(i)
                else:
                    console.print(f"Invalid number '{part}'.", style="error")
            except ValueError:
                console.print(f"Invalid input '{part}'.", style="error")
    return sorted(selections)

def parse_directory_input(directory_input):
    directory = directory_input.strip('"\'')
    if '/' in directory:
        directory = directory.rsplit('/', 1)[0]
    elif '\\' in directory:
        directory = directory.rsplit('\\', 1)[0]
    return directory

def read_recent_directories():
    recent_dirs = []
    if os.path.exists('recent_dirs.txt'):
        with open('recent_dirs.txt', 'r') as file:
            recent_dirs = [line.strip() for line in file.readlines()]
    return recent_dirs[:3]

def write_recent_directory(directory):
    recent_dirs = read_recent_directories()
    if directory in recent_dirs:
        recent_dirs.remove(directory)
    recent_dirs.insert(0, directory)
    with open('recent_dirs.txt', 'w') as file:
        for dir_ in recent_dirs[:3]:
            file.write(dir_ + '\n')

def recall_tables():
    """
    Menu system to let user pick a base directory (with recent shortcuts),
    then pick an ingested subfolder, then display CSV tables.
    """
    print("\033c", end="")
    title()
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")
    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(You can select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)
    if os.path.isdir(directory):
        write_recent_directory(directory)
        folders = list_processed_folders(directory)
        if not folders:
            return
        try:
            choice = int(console.input("Select a folder number to display tables (0 to cancel): "))
        except ValueError:
            console.print("Invalid input.", style="error")
            return
        if choice == 0:
            return
        if 1 <= choice <= len(folders):
            selected_folder = folders[choice - 1]
            output_directory = os.path.join(directory, selected_folder)
            csv_files = [f for f in os.listdir(output_directory) if f.endswith('.csv')]
            if not csv_files:
                console.print("No tables found in this folder.", style="warning")
                return
            console.print("Available tables in the folder:", style="header")
            for idx, csv_file in enumerate(csv_files, 1):
                console.print(f"{idx}. {csv_file}")
            while True:
                user_input = console.input("Enter table numbers to display (e.g., 1,2-5,8), or 0 to go back: ")
                if user_input.strip() == '0':
                    return
                selected_indices = parse_user_selection(user_input, len(csv_files))
                if selected_indices:
                    display_tables_from_folder(output_directory, selected_indices)
                    more_input = console.input("Would you like to view more tables from this folder? (y/n): ").strip().lower()
                    if more_input != 'y':
                        return
                else:
                    console.print("No valid table numbers entered. Please try again.", style="error")
        else:
            console.print("Invalid selection.", style="error")
    else:
        console.print("Invalid directory. Please try again.", style="error")

# ---------------------------------------------------------------------
# Explanation & Google Dork Functions
# ---------------------------------------------------------------------
def chunked_print(text, lines_per_page=30):
    """
    Prints the given text in chunks of `lines_per_page`.
    After printing each chunk, waits for user input:
      - Press Enter to continue
      - Type 'q' and press Enter to quit early
    """
    lines = text.strip().splitlines()
    idx = 0
    total_lines = len(lines)
    
    while idx < total_lines:
        chunk = lines[idx : idx + lines_per_page]
        # Print this chunk
        console.print("\n".join(chunk), style="info")
        
        idx += lines_per_page
        if idx < total_lines:
            # Prompt user if there's more to show
            user_input = console.input("\nPress [Enter] to continue or 'q' to quit: ").strip().lower()
            if user_input == 'q':
                break

def explain_tool():
    """
    Comprehensive explanation reflecting every key feature in the tool:
      - Document ingestion (IR.21, IR.85, etc.)
      - Table extraction and CSV saving
      - IP and CIDR extraction
      - ASN lookups
      - BGP queries (BGPView)
      - Resolved IP hostnames
      - Previously ingested data recall
      - OSINT/Recon display from BGP data
      - Google Dorking for new Telco docs
    """
    # Show a short header
    console.print("\n[+] What is this tool about?", style="header")

    # Build one multiline string to explain all features
    explanation_text = r"""
This tool is a one-stop solution to automate Open-Source Intelligence (OSINT) and
telecom-focused data extraction from common carrier documentation such as GSMA IR.21
and IR.85 files. Whether you are performing security research, network engineering,
or investigating roaming agreements, this suite of features will streamline your process.

Here is a breakdown of the major features:

1) Google Dorking for Telco Docs (IR.21 or IR.85)
   - Quickly discover publicly exposed IR.21/IR.85 PDFs using Google Dork queries.
   - Download interesting documents for further ingestion.

2) Ingest & Process New Documents (PDF or DOCX)
   - Automatically parse IR.21, IR.85, and similar telecom-related documents.
   - Detect and extract relevant tables containing:
       • Roaming details
       • SCCP Gateway info
       • DPC (Destination Point Code) data
       • IP addresses for Diameter, GPRS, LTE
       • ASN references, etc.
   - Save each extracted table as CSV for easy review.

3) Recall & Display Ingested Data
   - Organize extracted tables by folder for quick retrieval.
   - Quickly list previously ingested documents and pick which tables to display.
   - Use a Rich-based table view in the console for a clean, readable format.

4) Extract IPs & CIDRs
   - Parse your CSV data for any IP addresses or CIDR ranges.
   - Flatten or expand CIDRs to get all IPs if desired.
   - Save them in a unique list (unique_ips.txt) for further analysis or lookups.

5) ASN Lookup & BGP Queries (via BGPView)
   - Convert your extracted IPs/CIDRs into a minimal set of aggregated networks.
   - Query each network to identify the associated ASN (using IPWhois).
   - Optionally query BGPView for each ASN to see:
       • Prefixes (IPv4 & IPv6)
       • Peers, Upstreams, Downstreams
       • Internet Exchange Points (IXs)
   - Store this extended ASN data locally (asn_all_endpoints.json).

6) Resolve IPs to Hostnames
   - Perform reverse DNS lookups on IPs to find potential hostnames.
   - Handy for OSINT when identifying domain or service behind an IP.

7) OSINT & Recon Display
   - Provide a condensed, security-oriented view of BGP results.
   - Show relevant data (ASN info, contact emails, IP prefixes, peers, etc.) in Rich tables.
   - Great for threat research, network mapping, or cross-referencing telecom operators.

8) Misc. Utilities
   - Menu-based interface to walk through each feature step by step.
   - Persistent “recent directories” for quick navigation.
   - Color-coded console output powered by Rich.

In essence, this tool streamlines telecom OSINT from finding exposed IR.21/IR.85
documents online to extracting data, identifying IP addresses, performing ASN
analysis, and generating robust BGP-based intelligence. Its modular approach
lets you pick exactly what you want to do, whether it's table extraction, IP
analysis, or in-depth BGP recon.

"""

    # Now print this explanation in chunks of 30 lines
    chunked_print(explanation_text, lines_per_page=30)

def google_dork_telco_docs():
    """
    Sub-menu to pick IR.21 or IR.85 for the Google Dork query.
    """
    console.print("\n[+] Telco Document Google Dorking", style="header")
    console.print("Which document type would you like to search for?")
    console.print("1. IR.21")
    console.print("2. IR.85")
    choice = console.input("Enter your choice (1 or 2): ").strip()

    if choice == '1':
        query = "filetype:pdf IR.21"
        console.print("\nPerforming Google Dork for IR.21 docs...", style="info")
    elif choice == '2':
        query = "filetype:pdf IR.85"
        console.print("\nPerforming Google Dork for IR.85 docs...", style="info")
    else:
        console.print("[warning]Invalid choice. Returning to menu.[/warning]")
        return

    search_results = []
    try:
        from googlesearch import search
        for result in search(query, num_results=10):
            search_results.append(result)
    except Exception as e:
        console.print(f"[error]Error while performing Google search: {e}[/error]")
        return

    if not search_results:
        console.print("[warning]No results found.[/warning]")
        return

    console.print("\n[+] Found the following PDF links:", style="header")
    for idx, link in enumerate(search_results, 1):
        console.print(f"{idx}. {link}")

    while True:
        try:
            choice = int(console.input("\nEnter the number of the PDF you want to download (or 0 to cancel): ").strip())
            if choice == 0:
                return
            elif 1 <= choice <= len(search_results):
                selected_link = search_results[choice - 1]
                download_pdf(selected_link)
                return
            else:
                console.print("[error]Invalid selection. Please choose a valid number from the list.[/error]")
        except ValueError:
            console.print("[error]Invalid input. Please enter a number.[/error]")

def download_pdf(url):
    save_dir = "./my_ir_docs"
    os.makedirs(save_dir, exist_ok=True)
    file_name = url.split("/")[-1]
    file_path = os.path.join(save_dir, file_name)
    console.print(f"\n[+] Downloading PDF from {url}...", style="info")
    try:
        response = requests.get(url, stream=True)
        response.raise_for_status()
        with open(file_path, 'wb') as file:
            shutil.copyfileobj(response.raw, file)
        console.print(f"[success]Downloaded PDF to: {file_path}[/success]")
    except Exception as e:
        console.print(f"[error]Failed to download PDF: {e}[/error]")

# ---------------------------------------------------------------------
# IP & CIDR Extraction (only) + ASN & BGP (separate)
# ---------------------------------------------------------------------
BGPVIEW_BASE_URL = "https://api.bgpview.io"

def extract_ips_and_cidrs_from_row(row):
    extracted_ips = []
    for cell in row:
        cidr_matches = re.findall(r"\b\d{1,3}(?:\.\d{1,3}){3}/\d{1,2}\b", cell)
        extracted_ips.extend(cidr_matches)

        ip_matches = re.findall(r"\b\d{1,3}(?:\.\d{1,3}){3}\b", cell)
        for ip in ip_matches:
            try:
                ip_address(ip)
                extracted_ips.append(ip)
            except ValueError:
                continue
    return extracted_ips

def expand_cidrs_to_ips(ips_and_cidrs):
    expanded_ips = set()
    for item in ips_and_cidrs:
        try:
            if '/' in item:
                network = ip_network(item, strict=False)
                expanded_ips.update(str(ip) for ip in network)
            else:
                expanded_ips.add(item)
        except ValueError:
            console.print(f"[warning]Invalid IP or range skipped: {item}[/warning]")
    return sorted(expanded_ips)

def extract_ips_from_csv(directory):
    all_ips_and_cidrs = []
    csv_files = [f for f in os.listdir(directory) if f.endswith(".csv")]
    if not csv_files:
        console.print("[warning]No CSV files found in this folder.[/warning]")
        return all_ips_and_cidrs

    for filename in csv_files:
        file_path = os.path.join(directory, filename)
        console.print(f"Processing CSV: [info]{file_path}[/info]", style="info")
        with open(file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            _ = next(reader, None)  # skip header row if present
            for row in reader:
                extracted_ips = extract_ips_and_cidrs_from_row(row)
                if extracted_ips:
                    all_ips_and_cidrs.extend(extracted_ips)
    return all_ips_and_cidrs

def determine_cidrs(ips):
    ip_objs = []
    for ip_str in ips:
        try:
            ip_objs.append(ip_address(ip_str))
        except ValueError:
            pass
    collapsed_nets = collapse_addresses(ip_objs)
    return sorted(collapsed_nets, key=lambda n: (n.network_address, n.prefixlen))

def query_asn_for_cidr(ip_net):
    representative_ip = str(ip_net.network_address)
    try:
        obj = IPWhois(representative_ip)
        results = obj.lookup_rdap()
        return {
            "CIDR": str(ip_net),
            "Representative IP": representative_ip,
            "ASN": results.get("asn", "N/A"),
            "ASN Description": results.get("asn_description", "N/A"),
            "ASN Country": results.get("asn_country_code", "N/A")
        }
    except Exception as e:
        return {
            "CIDR": str(ip_net),
            "Representative IP": representative_ip,
            "ASN": None,
            "ASN Description": f"Error: {e}",
            "ASN Country": "N/A"
        }

def query_bgpview_all(asn, sleep_time=1):
    endpoints = [
        f"/asn/AS{asn}",
        f"/asn/AS{asn}/prefixes",
        f"/asn/AS{asn}/peers",
        f"/asn/AS{asn}/upstreams",
        f"/asn/AS{asn}/downstreams",
        f"/asn/AS{asn}/ixs",
    ]
    results = {}
    for endpoint in endpoints:
        url = BGPVIEW_BASE_URL + endpoint
        try:
            resp = requests.get(url, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            key = endpoint.lstrip('/')
            results[key] = data.get("data", {})
        except requests.exceptions.RequestException as e:
            results[endpoint] = {"error": str(e)}
        time.sleep(sleep_time)
    return results

def menu_extract_ips_only():
    """
    Just the IP extraction from previously ingested CSVs (no ASN or BGP).
    """
    print("\033c", end="")
    title()
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")
    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(You can select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return
    write_recent_directory(directory)

    folders = list_processed_folders(directory)
    if not folders:
        return
    try:
        choice = int(console.input("Select a folder number for IP extraction (0 to cancel): "))
    except ValueError:
        console.print("Invalid input.", style="error")
        return
    if choice == 0:
        return
    if 1 <= choice <= len(folders):
        selected_folder = folders[choice - 1]
        folder_path = os.path.join(directory, selected_folder)

        console.print(f"\n[info]Extracting IPs/CIDRs from CSVs in {folder_path}[/info]")
        all_ips_and_cidrs = extract_ips_from_csv(folder_path)
        if not all_ips_and_cidrs:
            console.print("[warning]No IPs or CIDRs found.[/warning]")
            return

        console.print("\nExpanding CIDRs into individual IPs...", style="info")
        unique_ips = expand_cidrs_to_ips(all_ips_and_cidrs)

        output_file_txt = os.path.join(folder_path, "unique_ips.txt")
        with open(output_file_txt, "w") as outfile:
            outfile.write("\n".join(unique_ips))
        console.print(f"[success]Unique IPs saved to: {output_file_txt}[/success]")

    else:
        console.print("Invalid selection.", style="error")

def menu_bgp_lookup_for_ips():
    """
    Separate function for ASN & BGP lookups from previously extracted IPs.
    """
    print("\033c", end="")
    title()
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")
    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(You can select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return
    write_recent_directory(directory)

    # Step: pick a subfolder
    folders = list_processed_folders(directory)
    if not folders:
        return
    try:
        choice = int(console.input("Select a folder to do ASN & BGP lookups (0 to cancel): "))
    except ValueError:
        console.print("Invalid input.", style="error")
        return
    if choice == 0:
        return
    if 1 <= choice <= len(folders):
        selected_folder = folders[choice - 1]
        folder_path = os.path.join(directory, selected_folder)

        # 1) Load the previously extracted IPs
        unique_ips_file = os.path.join(folder_path, "unique_ips.txt")
        if not os.path.isfile(unique_ips_file):
            console.print(f"[warning]No unique_ips.txt found in {folder_path}. Please extract IPs first.[/warning]")
            return

        with open(unique_ips_file, "r") as uf:
            unique_ips = [line.strip() for line in uf if line.strip()]
        if not unique_ips:
            console.print("[warning]The unique_ips.txt file is empty.[/warning]")
            return

        # 2) Collapsing IPs -> minimal set of CIDRs
        console.print("\nCollapsing IPs into minimal set of CIDRs...", style="info")
        collapsed_nets = determine_cidrs(unique_ips)

        # 3) Query ASN for each
        asn_results = []
        console.print("\nQuerying ASN info for each collapsed CIDR...", style="info")
        for net in collapsed_nets:
            asn_info = query_asn_for_cidr(net)
            asn_results.append(asn_info)
            console.print(
                f"CIDR: {asn_info['CIDR']}, ASN: {asn_info['ASN']}, Desc: {asn_info['ASN Description']}"
            )

        output_file_json = os.path.join(folder_path, "asn_results.json")
        with open(output_file_json, "w") as outfile:
            json.dump(asn_results, outfile, indent=4)
        console.print(f"[success]ASN results saved to: {output_file_json}[/success]")

        # 4) Ask user if they want BGPView data
        do_bgpview = console.input("\nWould you like to fetch all BGPView data for each ASN? (yes/no): ").strip().lower()
        if do_bgpview in ["yes", "y"]:
            asn_set = set()
            for entry in asn_results:
                asn_val = entry.get("ASN")
                if asn_val and asn_val.lower() != "none":
                    asn_val = str(asn_val).replace("AS", "")
                    asn_set.add(asn_val)

            console.print("\nQuerying BGPView for all sub-endpoints...\n", style="info")
            bgp_extended_data = {}
            for asn in sorted(asn_set):
                console.print(f"Pulling all BGPView data for ASN {asn} ...", style="info")
                bgp_data = query_bgpview_all(asn, sleep_time=1)
                bgp_extended_data[asn] = bgp_data

            all_endpoints_json = os.path.join(folder_path, "asn_all_endpoints.json")
            with open(all_endpoints_json, "w") as af:
                json.dump(bgp_extended_data, af, indent=4)
            console.print(f"[success]Detailed BGPView info saved to: {all_endpoints_json}[/success]")
        else:
            console.print("[info]Skipping BGPView data retrieval.[/info]")
    else:
        console.print("Invalid selection.", style="error")

def menu_resolve_ips():
    """
    Reads unique_ips.txt from a previously ingested folder,
    tries to resolve each IP to a hostname via socket,
    displays the results, and saves them to a file.
    """
    print("\033c", end="")
    title()
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")

    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(You can select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return
    write_recent_directory(directory)

    # Let user pick subfolder
    folders = list_processed_folders(directory)
    if not folders:
        return

    try:
        choice = int(console.input("Select a folder number for IP resolution (0 to cancel): "))
    except ValueError:
        console.print("Invalid input.", style="error")
        return
    if choice == 0:
        return

    if 1 <= choice <= len(folders):
        selected_folder = folders[choice - 1]
        folder_path = os.path.join(directory, selected_folder)

        # unique_ips.txt must exist
        unique_ips_file = os.path.join(folder_path, "unique_ips.txt")
        if not os.path.isfile(unique_ips_file):
            console.print(f"[warning]No unique_ips.txt found in {folder_path}. Please extract IPs first.[/warning]")
            return

        with open(unique_ips_file, "r") as uf:
            ips = [line.strip() for line in uf if line.strip()]

        if not ips:
            console.print("[warning]No IPs found in unique_ips.txt.[/warning]")
            return

        # Resolve each IP to hostname
        resolved_data = []
        console.print(f"\nResolving {len(ips)} IPs...", style="info")

        for ip_str in ips:
            hostname = None
            try:
                hostname = socket.gethostbyaddr(ip_str)[0]
            except socket.herror:
                hostname = "No reverse DNS entry"
            except Exception as e:
                hostname = f"Error: {e}"

            resolved_data.append({"IP": ip_str, "Hostname": hostname})
            console.print(f"IP: {ip_str} -> {hostname}")

        # Save to file
        output_file = os.path.join(folder_path, "resolved_ips.txt")
        with open(output_file, "w") as f:
            for entry in resolved_data:
                f.write(f"{entry['IP']},{entry['Hostname']}\n")

        console.print(f"[success]Resolved IPs saved to: {output_file}[/success]")

    else:
        console.print("Invalid selection.", style="error")



# ---------------------------------------------------------------------
# NEW: Display BGPView All Endpoints Menu
# ---------------------------------------------------------------------
def menu_display_bgp_endpoints():
    """
    Loads the asn_all_endpoints.json from a previously ingested folder
    (where BGP data was saved). Displays it in console, or optionally,
    show only a subset. This is a simple raw display for demonstration.
    """
    print("\033c", end="")
    title()
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")

    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(You can select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return
    write_recent_directory(directory)

    # Let user pick subfolder
    folders = list_processed_folders(directory)
    if not folders:
        return

    try:
        choice = int(console.input("Select a folder number to display BGP endpoints (0 to cancel): "))
    except ValueError:
        console.print("Invalid input.", style="error")
        return
    if choice == 0:
        return

    if 1 <= choice <= len(folders):
        selected_folder = folders[choice - 1]
        folder_path = os.path.join(directory, selected_folder)

        # asn_all_endpoints.json must exist
        bgp_file = os.path.join(folder_path, "asn_all_endpoints.json")
        if not os.path.isfile(bgp_file):
            console.print(f"[warning]No asn_all_endpoints.json found in {folder_path}. Please run BGP lookups first.[/warning]")
            return

        with open(bgp_file, "r") as bf:
            bgp_data = json.load(bf)

        # Display raw or partial
        console.print("\n[info]Found BGP data for the following ASNs:[/info]", style="header")
        for asn in bgp_data.keys():
            console.print(f"  ASN: {asn}")

        console.print("\nEnter an ASN to see detailed data, or 'all' to display everything, or 0 to cancel.")
        user_asn = console.input("ASN choice: ").strip().lower()
        if user_asn == '0':
            return

        if user_asn == 'all':
            console.print(json.dumps(bgp_data, indent=4), style="info")
        else:
            # If the user typed '12345' or 'AS12345'
            user_asn_str = user_asn.replace("as", "").strip()
            if user_asn_str in bgp_data:
                console.print(json.dumps(bgp_data[user_asn_str], indent=4), style="info")
            else:
                console.print(f"[warning]ASN {user_asn_str} not found in asn_all_endpoints.json[/warning]")
    else:
        console.print("Invalid selection.", style="error")


def display_osint_recon(bgp_data):
    """
    Displays interesting OSINT data from the BGP queries in a concise manner.
    Loops through each ASN in the JSON, gathers:
      - Basic ASN info
      - RIR/whois info
      - IPv4/IPv6 prefixes
      - Peers, Upstreams, Downstreams
      - IXs, etc.
    Also handles the case where asn/AS####/ixs might be a dict or a list.
    """
    for asn_str, endpoints_dict in bgp_data.items():
        console.rule(title=f"[bold magenta]ASN {asn_str}[/bold magenta]", style="header")
        
        # base_key: e.g. "asn/AS197465" or "asn/AS55083"
        base_key = f"asn/AS{asn_str}"
        
        # -------------------------
        # 1) Basic ASN Info
        # -------------------------
        if base_key in endpoints_dict:
            _display_asn_basic_info(endpoints_dict[base_key])
        else:
            console.print(f"[warning]No base ASN info found for ASN {asn_str}[/warning]")

        # -------------------------
        # 2) Prefixes
        # -------------------------
        prefixes_key = base_key + "/prefixes"
        if prefixes_key in endpoints_dict:
            _display_asn_prefixes(endpoints_dict[prefixes_key])
        else:
            console.print(f"[warning]No prefix info found for ASN {asn_str}[/warning]")

        # -------------------------
        # 3) Peers
        # -------------------------
        peers_key = base_key + "/peers"
        if peers_key in endpoints_dict:
            _display_asn_peers(endpoints_dict[peers_key], "Peers")
        else:
            console.print(f"[warning]No peer info found for ASN {asn_str}[/warning]")

        # -------------------------
        # 4) Upstreams
        # -------------------------
        ups_key = base_key + "/upstreams"
        if ups_key in endpoints_dict:
            _display_asn_peers(endpoints_dict[ups_key], "Upstreams")
        else:
            console.print(f"[warning]No upstream info found for ASN {asn_str}[/warning]")

        # -------------------------
        # 5) Downstreams
        # -------------------------
        downs_key = base_key + "/downstreams"
        if downs_key in endpoints_dict:
            _display_asn_peers(endpoints_dict[downs_key], "Downstreams")
        else:
            console.print(f"[warning]No downstream info found for ASN {asn_str}[/warning]")

        # -------------------------
        # 6) IXs
        # -------------------------
        ixs_key = base_key + "/ixs"
        if ixs_key in endpoints_dict:
            ixs_data = endpoints_dict[ixs_key]
            # If it's a dict, check if "ixs" is present
            if isinstance(ixs_data, dict):
                if "ixs" in ixs_data and isinstance(ixs_data["ixs"], list):
                    _display_asn_ixs(ixs_data)
                else:
                    console.print(f"[info]No IX data for ASN {asn_str}[/info]")
            # If it's already a list, wrap it in a dict to keep _display_asn_ixs happy
            elif isinstance(ixs_data, list):
                _display_asn_ixs({"ixs": ixs_data})
            else:
                console.print(f"[warning]Unexpected ixs data type for ASN {asn_str}: {type(ixs_data)}[/warning]")
        else:
            console.print(f"[info]No IX data for ASN {asn_str}[/info]")

        console.print()  # Blank line after each ASN


def _display_asn_basic_info(data):
    """
    Shows key OSINT fields: name, country_code, email contacts, abuse contacts, etc.
    """
    table = Table(title="Basic ASN Info", show_header=True, header_style="bold magenta", box=box.SQUARE)
    table.add_column("Field", style="header")
    table.add_column("Value")

    # Fields you find interesting for OSINT:
    interesting_fields = {
        "ASN": data.get("asn", "N/A"),
        "Name": data.get("name", "N/A"),
        "Short Description": data.get("description_short", "N/A"),
        "Country": data.get("country_code", "N/A"),
        "Emails": ", ".join(data.get("email_contacts", [])),
        "Abuse Contacts": ", ".join(data.get("abuse_contacts", [])),
        "Owner Address": ", ".join(data.get("owner_address", [])),
        "Date Updated": data.get("date_updated", "N/A"),
    }

    # RIR/WHOIS Info
    rir_info = data.get("rir_allocation", {})
    iana_info = data.get("iana_assignment", {})
    interesting_fields["RIR Name"] = rir_info.get("rir_name", "N/A")
    interesting_fields["RIR Country"] = rir_info.get("country_code", "N/A")
    interesting_fields["Allocated"] = rir_info.get("date_allocated", "N/A")
    interesting_fields["Whois Server"] = iana_info.get("whois_server", "N/A")

    # Add each row to the table
    for field_name, value in interesting_fields.items():
        table.add_row(field_name, str(value))

    console.print(table)


def _display_asn_prefixes(data):
    """
    Shows IPv4/IPv6 prefixes in separate tables, focusing on prefix, name, description, country_code.
    """
    ipv4_prefixes = data.get("ipv4_prefixes", [])
    ipv6_prefixes = data.get("ipv6_prefixes", [])

    # IPv4
    if ipv4_prefixes:
        table_v4 = Table(title="IPv4 Prefixes", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_v4.add_column("Prefix")
        table_v4.add_column("ROA")
        table_v4.add_column("Name", overflow="fold")
        table_v4.add_column("Description", overflow="fold")
        table_v4.add_column("Country", overflow="fold")

        for p in ipv4_prefixes:
            prefix = p.get("prefix", "N/A")
            roa = p.get("roa_status", "N/A")
            name = p.get("name", "N/A")
            desc = p.get("description", "N/A")
            ctry = p.get("country_code", "N/A")
            table_v4.add_row(prefix, roa, str(name), str(desc), ctry)

        console.print(table_v4)
    else:
        console.print("[warning]No IPv4 prefixes found[/warning]")

    # IPv6
    if ipv6_prefixes:
        table_v6 = Table(title="IPv6 Prefixes", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_v6.add_column("Prefix")
        table_v6.add_column("ROA")
        table_v6.add_column("Name", overflow="fold")
        table_v6.add_column("Description", overflow="fold")
        table_v6.add_column("Country", overflow="fold")

        for p in ipv6_prefixes:
            prefix = p.get("prefix", "N/A")
            roa = p.get("roa_status", "N/A")
            name = p.get("name", "N/A")
            desc = p.get("description", "N/A")
            ctry = p.get("country_code", "N/A")
            table_v6.add_row(prefix, roa, str(name), str(desc), ctry)

        console.print(table_v6)
    else:
        console.print("[warning]No IPv6 prefixes found[/warning]")


def _display_asn_peers(data, title_str):
    """
    General function to display IPv4 / IPv6 peers, upstreams, or downstreams.
    """
    ipv4_field = f"ipv4_{title_str.lower()}"
    ipv6_field = f"ipv6_{title_str.lower()}"

    ipv4_list = data.get(ipv4_field, [])
    ipv6_list = data.get(ipv6_field, [])

    # If the field doesn't match exactly, fallback to "ipv4_peers" / "ipv6_peers"
    if not ipv4_list and data.get("ipv4_peers"):
        ipv4_list = data.get("ipv4_peers", [])
    if not ipv6_list and data.get("ipv6_peers"):
        ipv6_list = data.get("ipv6_peers", [])

    # Display IPv4
    if ipv4_list:
        table_v4 = Table(title=f"IPv4 {title_str}", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_v4.add_column("Peer ASN")
        table_v4.add_column("Name", overflow="fold")
        table_v4.add_column("Description", overflow="fold")
        table_v4.add_column("Country")

        for peer in ipv4_list:
            pasn = str(peer.get("asn", "N/A"))
            pname = peer.get("name", "N/A")
            pdesc = peer.get("description", "N/A")
            pcc = peer.get("country_code", "N/A")
            table_v4.add_row(pasn, pname, pdesc, pcc)

        console.print(table_v4)
    else:
        console.print(f"[warning]No IPv4 {title_str.lower()} found[/warning]")

    # Display IPv6
    if ipv6_list:
        table_v6 = Table(title=f"IPv6 {title_str}", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_v6.add_column("Peer ASN")
        table_v6.add_column("Name", overflow="fold")
        table_v6.add_column("Description", overflow="fold")
        table_v6.add_column("Country")

        for peer in ipv6_list:
            pasn = str(peer.get("asn", "N/A"))
            pname = peer.get("name", "N/A")
            pdesc = peer.get("description", "N/A")
            pcc = peer.get("country_code", "N/A")
            table_v6.add_row(pasn, pname, pdesc, pcc)

        console.print(table_v6)
    else:
        console.print(f"[warning]No IPv6 {title_str.lower()} found[/warning]")


def _display_asn_ixs(data):
    """
    Display any Internet Exchange Points (IXs). Expects data to be a dict with a "ixs" key.
    If the user passes a list, they must wrap it in {"ixs": list}.
    """
    ixs_list = data.get("ixs", [])
    if not ixs_list:
        console.print("[warning]No IXS data found[/warning]")
        return

    table_ix = Table(title="Internet Exchange Points", show_header=True, header_style="bold magenta", box=box.SQUARE)
    table_ix.add_column("Name", overflow="fold")
    table_ix.add_column("City")
    table_ix.add_column("Country")
    table_ix.add_column("IPv4")
    table_ix.add_column("IPv6")

    for ix in ixs_list:
        name = ix.get("name", "N/A")
        city = ix.get("city", "N/A")
        country = ix.get("country", "N/A")
        ipv4 = ix.get("ipv4", "N/A")
        ipv6 = ix.get("ipv6", "N/A")
        table_ix.add_row(name, city, country, ipv4, ipv6)

    console.print(table_ix)

def menu_osint_recon():
    """
    Menu option to display OSINT/Recon data from the asn_all_endpoints.json file
    in previously ingested subfolders.
    """
    print("\033c", end="")
    title()  # or banner(), whichever you prefer
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")

    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(You can select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return
    write_recent_directory(directory)

    folders = list_processed_folders(directory)
    if not folders:
        return

    try:
        choice = int(console.input("Select a folder number to display OSINT Recon (0 to cancel): "))
    except ValueError:
        console.print("Invalid input.", style="error")
        return
    if choice == 0:
        return

    if 1 <= choice <= len(folders):
        selected_folder = folders[choice - 1]
        folder_path = os.path.join(directory, selected_folder)

        # Must have the asn_all_endpoints.json file
        bgp_file = os.path.join(folder_path, "asn_all_endpoints.json")
        if not os.path.isfile(bgp_file):
            console.print(f"[warning]No asn_all_endpoints.json found in {folder_path}. Please run BGP lookups first.[/warning]")
            return

        with open(bgp_file, "r") as bf:
            bgp_data = json.load(bf)

        # Now call our OSINT/Recon display function
        display_osint_recon(bgp_data)
    else:
        console.print("Invalid selection.", style="error")

def _display_dnsdumpster(domain, result_json):
    """
    Display the relevant data from DNSDumpster's JSON response in Rich tables.
    Example keys:
        "a" -> list of dicts for A records
        "mx" -> list of dicts for MX
        "ns" -> list of dicts for NS
        "txt" -> list of strings
        "cname" -> list of dicts for CNAME
        "total_a_recs" -> integer
    If map=1 was used, there may also be "domain_map" base64 data in the response.
    """
    # Show A records
    a_records = result_json.get("a", [])
    if a_records:
        table_a = Table(title=f"A Records for {domain}", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_a.add_column("Host")
        table_a.add_column("IP(s)")
        table_a.add_column("Country")
        table_a.add_column("ASN Range")

        for record in a_records:
            host = record.get("host", "N/A")
            # record["ips"] is typically a list of IP info dicts
            ip_string = []
            ctry_string = []
            asn_string = []
            for ip_obj in record.get("ips", []):
                ip_string.append(ip_obj.get("ip", "N/A"))
                ctry_string.append(ip_obj.get("country_code", "N/A"))
                asn_string.append(ip_obj.get("asn_range", "N/A"))
            # Join them if multiple
            table_a.add_row(
                host,
                "\n".join(ip_string),
                "\n".join(ctry_string),
                "\n".join(asn_string)
            )
        console.print(table_a)
    else:
        console.print("[info]No A records found.[/info]")

    # Show MX records
    mx_records = result_json.get("mx", [])
    if mx_records:
        table_mx = Table(title=f"MX Records for {domain}", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_mx.add_column("Host")
        table_mx.add_column("IP")
        table_mx.add_column("Country")
        table_mx.add_column("ASN Range")
        for mx in mx_records:
            host = mx.get("host", "N/A")
            ip_objs = mx.get("ips", [])
            for ip_obj in ip_objs:
                ip_ = ip_obj.get("ip", "N/A")
                ctry_ = ip_obj.get("country_code", "N/A")
                asn_ = ip_obj.get("asn_range", "N/A")
                table_mx.add_row(host, ip_, ctry_, asn_)
        console.print(table_mx)
    else:
        console.print("[info]No MX records found.[/info]")

    # Show NS records
    ns_records = result_json.get("ns", [])
    if ns_records:
        table_ns = Table(title=f"NS Records for {domain}", show_header=True, header_style="bold magenta", box=box.SQUARE)
        table_ns.add_column("Name Server")
        table_ns.add_column("IP")
        table_ns.add_column("Country")
        table_ns.add_column("ASN Range")
        for ns in ns_records:
            host = ns.get("host", "N/A")
            for ip_obj in ns.get("ips", []):
                ip_ = ip_obj.get("ip", "N/A")
                ctry_ = ip_obj.get("country_code", "N/A")
                asn_ = ip_obj.get("asn_range", "N/A")
                table_ns.add_row(host, ip_, ctry_, asn_)
        console.print(table_ns)
    else:
        console.print("[info]No NS records found.[/info]")

    # Show TXT
    txt_records = result_json.get("txt", [])
    if txt_records:
        console.print("\n[bold magenta]TXT Records:[/bold magenta]", style="info")
        for txt in txt_records:
            console.print(f"  {txt}", style="info")
    else:
        console.print("[info]No TXT records found.[/info]")

    # If the response had a domain map in base64, handle it
    # (Users with "Plus" membership might see "domain_map" in the JSON)
    domain_map = result_json.get("domain_map", None)
    if domain_map:
        console.print("\n[success]A domain map was included (base64).[/success]")
        # Optionally save it to a file:
        out_map = f"{domain.replace('.', '_')}_map.png"
        try:
            with open(out_map, "wb") as mf:
                mf.write(base64.b64decode(domain_map))
            console.print(f"[success]Domain map saved to {out_map}[/success]")
        except Exception as e:
            console.print(f"[warning]Could not decode/save domain map: {e}[/warning]")

def dnsdumpster_enum(domain, api_key, page="1", map_image=False):
    """
    Performs a DNSDumpster API query for the given domain, returning parsed JSON (dict)
    or None if there's an error.
    """
    base_url = f"https://api.dnsdumpster.com/domain/{domain}"
    params = {}

    # Pagination: ?page=2
    if page not in ("", "1"):
        params["page"] = page

    # Map: ?map=1
    if map_image:
        params["map"] = "1"

    headers = {
        "X-API-Key": api_key
    }

    try:
        time.sleep(2)  # DNSDumpster rate limit is 1 request / 2 seconds
        response = requests.get(base_url, headers=headers, params=params, timeout=30)
        response.raise_for_status()  # Raise if HTTP 4xx/5xx
        # Example of rate limit exceeded => 429 => raise_for_status triggers an exception

        data = response.json()
        return data

    except requests.exceptions.HTTPError as http_err:
        console.print(f"[error]HTTP error occurred: {http_err}[/error]")
    except requests.exceptions.RequestException as e:
        console.print(f"[error]Request error occurred: {e}[/error]")
    except Exception as e:
        console.print(f"[error]Unexpected error: {e}[/error]")

    return None  # In case of error

def menu_subdomain_dnsdumpster():
    """
    Menu option to perform subdomain enumeration using DNSDumpster's API.
    Requires an API key from dnsdumpster.com.
    """
    print("\033c", end="")
    title()
    console.print("[+] DNSDumpster Subdomain Enumeration", style="header")

    # Ask for domain
    domain = console.input("Enter the domain you want to query (e.g., example.com): ").strip()
    if not domain:
        console.print("[error]No domain entered. Returning to main menu.[/error]")
        return

    # Ask for API key
    api_key = console.input("Enter your DNSDumpster API key: ").strip()
    if not api_key:
        console.print("[error]No API key entered. Returning to main menu.[/error]")
        return

    # Optional parameters
    console.print("Enter page number to request (e.g. 1, 2, 3...), or press Enter for page=1:", style="info")
    page_input = console.input().strip()
    page = page_input if page_input else "1"

    console.print("Include domain map? (y/n) [Plus membership required for map]", style="info")
    map_input = console.input().strip().lower()
    map_image = (map_input == 'y')

    # Make the request
    console.print(f"\nQuerying DNSDumpster for [bold cyan]{domain}[/bold cyan], page={page}, map={map_image}", style="info")
    result_json = dnsdumpster_enum(domain, api_key, page, map_image)
    if not result_json:
        console.print("[warning]No data returned or error occurred.[/warning]")
        return

    # Parse and display results
    console.print("[success]Successfully retrieved DNSDumpster results![/success]")
    _display_dnsdumpster(domain, result_json)

    # Save to file
    out_filename = f"dnsdumpster_{domain.replace('.', '_')}_page{page}.json"
    with open(out_filename, "w") as outfile:
        json.dump(result_json, outfile, indent=4)
    console.print(f"[success]Data saved to {out_filename}[/success]")

def menu_subdomain_enum():
    """
    Menu option to perform subdomain enumeration on a user-specified domain.
    """
    print("\033c", end="")
    title()
    console.print("[+] Subdomain Enumeration", style="header")

    # 1) Ask for the domain name
    domain = console.input("Enter the domain you want to enumerate subdomains for (e.g., example.com): ").strip()
    if not domain:
        console.print("[error]No domain entered. Returning to main menu.[/error]")
        return

    # 2) Optionally let user specify a wordlist or default
    default_wordlist = "subdomains.txt"
    console.print(f"Enter path to subdomain wordlist, or press Enter to use default ({default_wordlist}):", style="info")
    wordlist_path = console.input().strip()
    if not wordlist_path:
        wordlist_path = default_wordlist

    # 3) Run enumeration
    console.print(f"\nEnumerating subdomains of [bold cyan]{domain}[/bold cyan] with wordlist [info]{wordlist_path}[/info] ...", style="info")
    found = enumerate_subdomains(domain, wordlist_path)

    # 4) Display results in a Rich table
    if found:
        console.print(f"\n[success]Found {len(found)} subdomains for {domain}[/success]")
        _display_subdomains(domain, found)
    else:
        console.print(f"[warning]No valid subdomains found for {domain}[/warning]")

    # 5) Save results to a file
    output_file = f"subdomains_{domain.replace('.', '_')}.txt"
    with open(output_file, "w") as f:
        for sub in found:
            f.write(sub + "\n")
    console.print(f"\n[success]Subdomains saved to: {output_file}[/success]")

def enumerate_subdomains(domain, wordlist_path="subdomains.txt"):
    """
    Enumerate subdomains by reading 'wordlist_path' and checking if
    subdomain.domain resolves to an IP address.
    Returns a list of valid subdomains.
    """
    valid_subdomains = []
    try:
        with open(wordlist_path, "r") as wf:
            sub_names = [line.strip() for line in wf if line.strip()]
    except FileNotFoundError:
        console.print(f"[error]Could not open wordlist: {wordlist_path}[/error]")
        return valid_subdomains

    def check_sub(sub):
        full_domain = f"{sub}.{domain}"
        try:
            # Attempt DNS resolution
            socket.gethostbyname(full_domain)
            return full_domain
        except socket.gaierror:
            return None

    # Use a ThreadPoolExecutor to check subdomains in parallel
    with ThreadPoolExecutor(max_workers=20) as executor:
        future_to_sub = {executor.submit(check_sub, s): s for s in sub_names}
        for future in as_completed(future_to_sub):
            result = future.result()
            if result:
                valid_subdomains.append(result)

    # Sort results for neatness
    valid_subdomains.sort()
    return valid_subdomains

def _display_subdomains(domain, found_list):
    table = Table(title=f"Subdomains for {domain}", show_header=True, header_style="bold magenta", box=box.SQUARE)
    table.add_column("Index", justify="right")
    table.add_column("Subdomain", overflow="fold")

    for idx, sub in enumerate(found_list, start=1):
        table.add_row(str(idx), sub)

    console.print(table)

def menu_subdomain_dnsdumpster_folder():
    """
    Pull the domain(s) from resolved_ips.txt in one of the previously ingested folders
    and query DNSDumpster for them. Allows picking from multiple domains if found.
    """
    print("\033c", end="")
    title()
    console.print("[+] DNSDumpster Subdomain Enumeration from resolved_ips.txt", style="header")

    # 1) Pick a base directory (using your recent_dirs logic)
    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")

    console.print("Enter the base directory where ingested folders are stored:")
    base_dir_input = console.input("(Select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return
    write_recent_directory(directory)

    # 2) Let user pick subfolder
    folders = list_processed_folders(directory)
    if not folders:
        return

    try:
        choice = int(console.input("Select a folder to load resolved_ips.txt (0 to cancel): "))
    except ValueError:
        console.print("Invalid input.", style="error")
        return
    if choice == 0:
        return

    if 1 <= choice <= len(folders):
        selected_folder = folders[choice - 1]
        folder_path = os.path.join(directory, selected_folder)

        # 3) Read resolved_ips.txt
        resolved_file = os.path.join(folder_path, "resolved_ips.txt")
        if not os.path.isfile(resolved_file):
            console.print(f"[warning]No resolved_ips.txt found in {folder_path}. Please resolve IPs first.[/warning]")
            return

        # 4) Parse domains from the file
        domains_found = parse_domains_from_resolved(resolved_file)
        if not domains_found:
            console.print("[warning]No valid domains found in resolved_ips.txt[/warning]")
            return

        # 5) If multiple domains, let the user pick
        chosen_domain = None
        if len(domains_found) == 1:
            chosen_domain = next(iter(domains_found))
            console.print(f"Found single domain: [bold cyan]{chosen_domain}[/bold cyan]")
        else:
            console.print("Found multiple domains:", style="header")
            domain_list = sorted(domains_found)
            for idx, d in enumerate(domain_list, 1):
                console.print(f"{idx}. {d}", style="info")

            try:
                domain_choice = int(console.input("Select domain to query (0 to cancel): ").strip())
                if domain_choice == 0:
                    return
                if 1 <= domain_choice <= len(domain_list):
                    chosen_domain = domain_list[domain_choice - 1]
                else:
                    console.print("Invalid selection.", style="error")
                    return
            except ValueError:
                console.print("Invalid input.", style="error")
                return

        # 6) Ask for API key, page, map param
        if chosen_domain:
            console.print(f"\nPreparing to query DNSDumpster for domain: [bold cyan]{chosen_domain}[/bold cyan]")
            api_key = console.input("Enter your DNSDumpster API key: ").strip()
            if not api_key:
                console.print("[error]No API key provided. Returning to menu.[/error]")
                return

            # Optional: ask for page, map
            console.print("Enter page number to request (e.g. 1, 2, 3...), or press Enter for page=1:", style="info")
            page_input = console.input().strip()
            page = page_input if page_input else "1"

            console.print("Include domain map? (y/n) [Plus membership required for map]", style="info")
            map_input = console.input().strip().lower()
            map_image = (map_input == 'y')

            # 7) Query DNSDumpster
            result_json = dnsdumpster_enum(chosen_domain, api_key, page, map_image)
            if not result_json:
                console.print("[warning]No data returned or an error occurred.[/warning]")
                return

            console.print("[success]Successfully retrieved DNSDumpster results![/success]")
            _display_dnsdumpster(chosen_domain, result_json)

            # Save JSON
            out_filename = os.path.join(folder_path, f"dnsdumpster_{chosen_domain.replace('.', '_')}_page{page}.json")
            with open(out_filename, "w") as outfile:
                json.dump(result_json, outfile, indent=4)
            console.print(f"[success]Data saved to {out_filename}[/success]")

    else:
        console.print("Invalid selection.", style="error")

def parse_domains_from_resolved(resolved_file):
    """
    Reads the lines of resolved_ips.txt, extracting domain names 
    from the 'hostname' column if present. 
    Returns a set of unique domains.
    """
    domains = set()
    with open(resolved_file, "r") as rf:
        for line in rf:
            line = line.strip()
            if not line:
                continue
            # "IP,hostname"
            parts = line.split(",")
            if len(parts) < 2:
                continue
            hostname = parts[1].strip()

            # Skip lines like "No reverse DNS entry"
            if not hostname or hostname.lower() == "no reverse dns entry":
                continue

            # Attempt a naive domain extraction from the hostname
            domain = extract_domain(hostname)
            if domain:
                domains.add(domain)
    return domains

def extract_domain(hostname):
    """
    Given a hostname like 'msn-nmrx-185-187.msn01.numerex.com',
    return 'numerex.com' by taking the last 2 components.
    If that doesn't work, return None.
    """
    parts = hostname.split(".")
    if len(parts) >= 2:
        return ".".join(parts[-2:])  # e.g. ['numerex','com'] => 'numerex.com'
    return None

# ---------------------------------------------------------------------
# IR.21 Keywords
# ---------------------------------------------------------------------
IR21_KEYWORDS = {
    "Basic Operator Identifiers": [
        "MCC", "MNC", "Mobile Country Code", "Mobile Network Code", "TADIG",
    ],
    "Numbering Plans & IMSI/ MSISDN Ranges": [
        "IMSI range", "MSISDN range", "Number series", "Subscriber number",
    ],
    "Roaming Agreements": [
        "Inbound Roaming", "Outbound Roaming", "Roaming Partners",
    ],
    "Signaling/Core Network Details": [
        "SCCP", "DPC", "MAP version", "CAP phase", "MSC", "VLR",
    ],
    "Diameter & LTE/EPC Data": [
        "Diameter", "S6a", "S6d", "IPX", "GRX", "Realm",
    ],
    "GPRS & Packet Core Info": [
        "GGSN", "SGSN", "APN", "DNS IP",
    ],
    "IP & CIDR Ranges": [
        "IP Address", "IPv4", "IPv6", "CIDR", "ASN",
    ],
    "Voice/Messaging": [
        "MSC", "SMSC", "MMS", "WAP",
    ],
    "Technical Contact": [
        "Roaming Contact", "Email", "Engineering Contact",
    ],
    "Legacy & Special Config": [
        "Circuit-Switched Fallback", "2G", "3G", "Partial CAMEL",
    ],
}

# ---------------------------------------------------------------------
# Helper Functions for PDF/DOCX Text & Table Extraction
# ---------------------------------------------------------------------
def extract_text_from_pdf(pdf_path):
    """Extract all text from the PDF for broad searching."""
    full_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            full_text.append(text)
    return "\n".join(full_text)

def extract_text_from_docx(docx_path):
    """Extract all text from the DOCX for broad searching."""
    doc = docx.Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)

def extract_tables_from_pdf(pdf_path):
    """Return a list of tables (each table is a list of rows; each row a list of strings)."""
    all_tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_tables = page.extract_tables()
            for table in page_tables or []:
                sanitized_table = []
                for row in table:
                    row = [str(cell) if cell else "" for cell in row]
                    sanitized_table.append(row)
                all_tables.append(sanitized_table)
    return all_tables

def extract_tables_from_docx(docx_path):
    """Return a list of tables (each table is a list of rows; each row a list of strings)."""
    doc = docx.Document(docx_path)
    all_tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        all_tables.append(data)
    return all_tables

# ---------------------------------------------------------------------
# Search for IR.21 fields in extracted text/tables
# ---------------------------------------------------------------------
def parse_ir21_fields(text, tables):
    """
    Given text and tables from an IR.21 PDF/DOCX, extract relevant fields
    and return them in a structured dict keyed by category.
    """
    findings = {cat: [] for cat in IR21_KEYWORDS}

    # 1) Search the raw text for matches
    for category, keywords in IR21_KEYWORDS.items():
        for kw in keywords:
            if re.search(rf"\b{re.escape(kw)}\b", text, re.IGNORECASE):
                findings[category].append(f"Found keyword in text: {kw}")

    # 2) Search inside each table cell
    for table in tables:
        for row in table:
            for cell in row:
                cell_lower = cell.lower()
                for category, keywords in IR21_KEYWORDS.items():
                    for kw in keywords:
                        if kw.lower() in cell_lower:
                            findings[category].append(f"Found '{kw}' in table cell: {row}")

    return findings

# ---------------------------------------------------------------------
# Saving and Displaying (Single Doc) IR.21 Data
# ---------------------------------------------------------------------
def save_ir21_data_as_json(ir21_data, output_file):
    """
    Save extracted IR.21 data (just for one doc) to JSON.
    ir21_data typically is the dictionary:
      {
        "Basic Operator Identifiers": [...],
        "Numbering Plans & IMSI/ MSISDN Ranges": [...],
        ...
      }
    """
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(ir21_data, f, indent=4)

def display_single_ir21_data_in_console(doc_name, doc_findings):
    """
    Display a single document's IR.21 data in a more human-readable manner.
    doc_findings is a dict of categories -> list of found strings.
    """
    console.print(f"\n[bold magenta]Document: {doc_name}[/bold magenta]\n")
    for category, findings in doc_findings.items():
        if not findings:
            continue
        # A simple Rich table for each category
        table = Table(title=f"{category}", box=box.MINIMAL, show_header=False)
        for item in findings:
            table.add_row(item)
        console.print(table)

# ---------------------------------------------------------------------
# Main function to scan a directory for IR.21 docs, extract & display
# ---------------------------------------------------------------------
def analyze_ir21_in_directory(directory):
    """
    - Finds PDF/DOCX in `directory`
    - Extracts text + tables
    - Parses IR.21 features
    - Saves each doc's data as DOCNAME_ir21.json
    - Displays each doc's results in a more readable format
    """
    if not os.path.isdir(directory):
        console.print(f"[red]Invalid directory: {directory}[/red]")
        return

    # We'll track if we found any docs
    processed_any = False

    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        if not os.path.isfile(filepath):
            continue

        if filename.lower().endswith(".pdf"):
            console.print(f"[green]Processing PDF:[/green] {filename}")
            text = extract_text_from_pdf(filepath)
            tables = extract_tables_from_pdf(filepath)

        elif filename.lower().endswith(".docx"):
            console.print(f"[green]Processing DOCX:[/green] {filename}")
            text = extract_text_from_docx(filepath)
            tables = extract_tables_from_docx(filepath)

        else:
            # Skip non-PDF/DOCX
            continue

        # Mark that we processed at least one doc
        processed_any = True

        findings = parse_ir21_fields(text, tables)

        # Save this doc's IR.21 data to {docbase}_ir21.json
        doc_base = os.path.splitext(filename)[0]
        output_json = os.path.join(directory, f"{doc_base}_ir21.json")
        save_ir21_data_as_json(findings, output_json)
        console.print(f"[bold yellow]IR.21 data saved to {output_json}[/bold yellow]")

        # Now display to console in a more user-friendly manner
        display_single_ir21_data_in_console(filename, findings)

    if not processed_any:
        console.print("[yellow]No PDF or DOCX files found in this folder, or no IR.21 patterns detected.[/yellow]")

# ---------------------------------------------------------------------
# Menu Option to Analyze IR.21
# ---------------------------------------------------------------------
def menu_analyze_ir21():
    """
    Menu option to scan a directory (or a previously ingested folder) for IR.21 documents,
    extract relevant data, then save and display results.
    """
    print("\033c", end="")
    title()
    console.print("[+] Analyze IR.21 Documents", style="header")

    recent_dirs = read_recent_directories()
    if recent_dirs:
        console.print("Recent directories:", style="header")
        for idx, dir_ in enumerate(recent_dirs, 1):
            console.print(f"{idx}. {dir_}")

    console.print("Enter the base directory where your IR.21 docs are stored:")
    base_dir_input = console.input("(Select a recent directory by number, or enter a new path): ").strip()
    if base_dir_input.isdigit():
        selection = int(base_dir_input)
        if 1 <= selection <= len(recent_dirs):
            directory = recent_dirs[selection - 1]
        else:
            console.print("Invalid selection.", style="error")
            return
    else:
        directory = parse_directory_input(base_dir_input)

    if not os.path.isdir(directory):
        console.print(f"[error]Invalid directory: {directory}[/error]")
        return

    write_recent_directory(directory)

    # Now analyze IR.21 data from each PDF/DOCX
    analyze_ir21_in_directory(directory)
    console.print("[bold green]IR.21 Analysis Complete.[/bold green]")



# ---------------------------------------------------------------------
# Main Menu - Updated
# ---------------------------------------------------------------------
def main_menu():
    banner()
    title()
    console.print("\n[+] Welcome to All-the-Things!", style="header")

    while True:
        console.print("\nPlease choose an option:")
        console.print("1. What does this tool do?")
        console.print("2. Google Dork for Telco Docs (IR.21 or IR.85)")
        console.print("3. Process New Documents")
        console.print("4. Analyze IR.21 Documents")
        console.print("5. Recall Tables from Previously Ingested Folders")
        console.print("6. Extract IP from Previously Ingested Docs")
        console.print("7. BGP Lookups of Previously Extracted IPs")
        console.print("8. Resolve Extracted IPs to Hostnames")
        console.print("9. Display Raw BGP Data from File")
        console.print("10. Parsed BGP Data")
        console.print("11. Subdomain Enumeration (Wordlist-based)")
        console.print("12. DNSDumpster Subdomain Enumeration from resolved_ips.txt")
        console.print("13. Exit")

        choice = console.input("Enter your choice (1-13): ").strip()

        if choice == '1':
            explain_tool()
        elif choice == '2':
            google_dork_telco_docs()
        elif choice == '3':
            directory_input = console.input("Enter the directory path containing the documents: ").strip()
            directory = parse_directory_input(directory_input)
            if os.path.isdir(directory):
                files = [os.path.join(directory, f) for f in os.listdir(directory)]
                for file_path in files:
                    process_file(file_path, directory)  # Your existing logic for new docs
                write_recent_directory(directory)
            else:
                console.print("Invalid directory. Please try again.", style="error")

        elif choice == '4':
            menu_analyze_ir21()

        elif choice == '5':
            recall_tables()

        elif choice == '6':
            menu_extract_ips_only()

        elif choice == '7':
            menu_bgp_lookup_for_ips()

        elif choice == '8':
            menu_resolve_ips()

        elif choice == '9':
            menu_display_bgp_endpoints()

        elif choice == '10':
            menu_osint_recon()

        elif choice == '11':
            menu_subdomain_enum()

        elif choice == '12':
            menu_subdomain_dnsdumpster_folder()

        elif choice == '13':
            console.print("Goodbye!", style="success")
            sys.exit(0)

        else:
            console.print("Invalid choice. Please try again.", style="error")

# ---------------------------------------------------------------------
# Demo usage if you run this standalone:
# ---------------------------------------------------------------------
if __name__ == "__main__":
    main_menu()